import fitz  # PyMuPDF
import os
import json
import re
import uuid
import zipfile
from PIL import Image
import io

# 尝试导入python-docx，如果失败则使用备用方案
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Word conversion will be disabled.")

def pdf_to_markdown(pdf_path):
    """
    将PDF文件转换为格式化的Markdown文本，包含图片提取
    
    参数:
        pdf_path (str): PDF文件的路径
        
    返回:
        dict: 包含markdown内容和图片信息的字典
    """
    try:
        # 验证文件路径
        print(f"尝试打开PDF文件: {pdf_path}")
        print(f"文件是否存在: {os.path.exists(pdf_path)}")
        print(f"当前工作目录: {os.getcwd()}")
        
        if not os.path.exists(pdf_path):
            return {
                "markdown_content": f"文件不存在: {pdf_path}",
                "images": []
            }
        
        # 打开PDF文件
        doc = fitz.open(pdf_path)
        
        if doc.page_count == 0:
            return {
                "markdown_content": "未找到内容。PDF文件为空。",
                "images": []
            }
        
        # 创建图片存储目录
        base_filename = os.path.splitext(os.path.basename(pdf_path))[0]
        images_dir = os.path.join(os.path.dirname(pdf_path), f"{base_filename}_images")
        if not os.path.exists(images_dir):
            os.makedirs(images_dir)
        
        # 提取文本块和格式信息
        text_elements = []
        extracted_images = []
        
        for page_num in range(doc.page_count):
            page = doc[page_num]
            
            # 提取图片
            page_images = extract_images_from_page(page, page_num, images_dir, base_filename)
            extracted_images.extend(page_images)
            
            # 获取文本字典，包含格式信息
            text_dict = page.get_text("dict")
            
            # 处理每个文本块
            for block in text_dict["blocks"]:
                if "lines" in block:  # 文本块
                    for line in block["lines"]:
                        for span in line["spans"]:
                            if span["text"].strip():
                                text_elements.append({
                                    "text": span["text"],
                                    "font_size": span["size"],
                                    "font_flags": span["flags"],
                                    "bbox": span["bbox"],
                                    "page": page_num + 1
                                })
        
        # 关闭文档
        doc.close()
        
        if not text_elements and not extracted_images:
            return {
                "markdown_content": "未找到文本或图片内容。",
                "images": []
            }
        
        # 处理文本元素，转换为Markdown
        markdown_text = convert_elements_to_markdown(text_elements, extracted_images)
        
        return {
            "markdown_content": markdown_text,
            "images": extracted_images
        }
    
    except Exception as e:
        # 记录错误并返回错误消息
        error_msg = f"转换过程中出错: {str(e)}"
        print(error_msg)
        return {
            "markdown_content": error_msg,
            "images": []
        }

def extract_images_from_page(page, page_num, images_dir, base_filename):
    """
    从PDF页面提取图片
    
    参数:
        page: PDF页面对象
        page_num: 页面编号
        images_dir: 图片保存目录
        base_filename: 基础文件名
        
    返回:
        list: 提取的图片信息列表
    """
    images = []
    image_list = page.get_images()
    
    for img_index, img in enumerate(image_list):
        try:
            # 获取图片数据
            xref = img[0]
            pix = fitz.Pixmap(page.parent, xref)
            
            # 跳过CMYK图片（转换复杂）
            if pix.n - pix.alpha < 4:
                # 生成唯一的图片文件名
                img_filename = f"{base_filename}_page{page_num + 1}_img{img_index + 1}.png"
                img_path = os.path.join(images_dir, img_filename)
                
                # 保存图片
                if pix.alpha:
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                
                pix.save(img_path)
                
                # 记录图片信息
                images.append({
                    "filename": img_filename,
                    "path": img_path,
                    "page": page_num + 1,
                    "index": img_index + 1,
                    "width": pix.width,
                    "height": pix.height
                })
                
            pix = None  # 释放内存
            
        except Exception as e:
            print(f"提取第{page_num + 1}页第{img_index + 1}张图片时出错: {str(e)}")
            continue
    
    return images

def convert_elements_to_markdown(text_elements, extracted_images):
    """
    将文本元素转换为格式化的Markdown，包含图片引用
    
    参数:
        text_elements (list): 包含格式信息的文本元素列表
        extracted_images (list): 提取的图片信息列表
        
    返回:
        str: 格式化的Markdown文本
    """
    if not text_elements and not extracted_images:
        return ""
    
    # 如果只有图片没有文本
    if not text_elements:
        markdown_lines = []
        for img in extracted_images:
            markdown_lines.append(f"![图片 {img['index']}]({img['filename']})")
            markdown_lines.append(f"*第{img['page']}页 - 图片{img['index']} ({img['width']}x{img['height']})*")
        return "\n\n".join(markdown_lines)
    
    # 分析字体大小，确定标题级别
    font_sizes = [elem["font_size"] for elem in text_elements]
    avg_font_size = sum(font_sizes) / len(font_sizes)
    max_font_size = max(font_sizes)
    
    # 定义标题阈值
    title_threshold = avg_font_size * 1.2
    subtitle_threshold = avg_font_size * 1.1
    
    markdown_lines = []
    current_paragraph = []
    last_y = None
    current_page = 1
    
    # 按页面组织图片
    images_by_page = {}
    for img in extracted_images:
        page = img['page']
        if page not in images_by_page:
            images_by_page[page] = []
        images_by_page[page].append(img)
    
    for i, elem in enumerate(text_elements):
        text = elem["text"].strip()
        font_size = elem["font_size"]
        font_flags = elem["font_flags"]
        bbox = elem["bbox"]
        page = elem["page"]
        
        if not text:
            continue
        
        # 检查是否切换到新页面，如果是则插入该页面的图片
        if page != current_page:
            # 完成当前段落
            if current_paragraph:
                paragraph_text = " ".join(current_paragraph).strip()
                if paragraph_text:
                    markdown_lines.append(paragraph_text)
                current_paragraph = []
            
            # 插入前一页的图片（如果有）
            if current_page in images_by_page:
                markdown_lines.append("")  # 添加空行
                for img in images_by_page[current_page]:
                    markdown_lines.append(f"![图片 {img['index']}]({img['filename']})")
                    markdown_lines.append(f"*第{img['page']}页 - 图片{img['index']} ({img['width']}x{img['height']})*")
                markdown_lines.append("")  # 添加空行
            
            current_page = page
            last_y = None  # 重置Y坐标
        
        # 检测换行（基于y坐标变化）
        current_y = bbox[1]  # y0坐标
        line_break = False
        
        if last_y is not None:
            y_diff = abs(current_y - last_y)
            if y_diff > font_size * 0.8:  # 行间距大于字体大小的80%
                line_break = True
        
        # 处理当前段落
        if line_break and current_paragraph:
            # 完成当前段落
            paragraph_text = " ".join(current_paragraph).strip()
            if paragraph_text:
                markdown_lines.append(paragraph_text)
            current_paragraph = []
        
        # 判断文本类型并格式化
        formatted_text = format_text_element(text, font_size, font_flags, 
                                           title_threshold, subtitle_threshold, avg_font_size)
        
        # 如果是标题，直接添加到结果中
        if formatted_text.startswith('#'):
            if current_paragraph:
                paragraph_text = " ".join(current_paragraph).strip()
                if paragraph_text:
                    markdown_lines.append(paragraph_text)
                current_paragraph = []
            markdown_lines.append(formatted_text)
        else:
            # 添加到当前段落
            current_paragraph.append(formatted_text)
        
        last_y = current_y
    
    # 处理最后一个段落
    if current_paragraph:
        paragraph_text = " ".join(current_paragraph).strip()
        if paragraph_text:
            markdown_lines.append(paragraph_text)
    
    # 插入最后一页的图片（如果有）
    if current_page in images_by_page:
        markdown_lines.append("")  # 添加空行
        for img in images_by_page[current_page]:
            markdown_lines.append(f"![图片 {img['index']}]({img['filename']})")
            markdown_lines.append(f"*第{img['page']}页 - 图片{img['index']} ({img['width']}x{img['height']})*")
    
    # 后处理：清理和优化
    processed_lines = post_process_markdown(markdown_lines)
    
    return "\n\n".join(processed_lines)

def format_text_element(text, font_size, font_flags, title_threshold, subtitle_threshold, avg_font_size):
    """
    根据字体信息格式化文本元素
    """
    # 清理文本
    text = re.sub(r'\s+', ' ', text.strip())
    
    # 检测标题
    if font_size >= title_threshold:
        if font_size >= title_threshold * 1.3:
            return f"# {text}"
        elif font_size >= title_threshold * 1.1:
            return f"## {text}"
        else:
            return f"### {text}"
    elif font_size >= subtitle_threshold:
        return f"#### {text}"
    
    # 检测粗体 (font_flags & 16 表示粗体)
    if font_flags & 16:
        return f"**{text}**"
    
    # 检测斜体 (font_flags & 2 表示斜体)
    if font_flags & 2:
        return f"*{text}*"
    
    # 检测列表项
    if re.match(r'^[\d]+\.', text):  # 数字列表
        return text
    elif re.match(r'^[•·▪▫‣⁃]\s*', text):  # 项目符号列表
        return re.sub(r'^[•·▪▫‣⁃]\s*', '- ', text)
    elif re.match(r'^[-*+]\s+', text):  # 已经是Markdown列表格式
        return text
    
    return text

def post_process_markdown(lines):
    """
    后处理Markdown文本，优化格式
    """
    processed = []
    
    for i, line in enumerate(lines):
        if not line.strip():
            continue
        
        # 处理列表项
        if re.match(r'^[\d]+\.', line) or line.startswith('- '):
            # 确保列表项前后有适当的空行
            if processed and not processed[-1].startswith(('- ', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '0.')):
                if not processed[-1].startswith('#'):
                    processed.append('')  # 在列表前添加空行
            processed.append(line)
        else:
            processed.append(line)
    
    # 移除多余的空行
    final_lines = []
    for i, line in enumerate(processed):
        if line == '' and i > 0 and processed[i-1] == '':
            continue  # 跳过连续的空行
        final_lines.append(line)
    
    return final_lines

def pdf_to_word(pdf_path, output_path=None):
    """
    将PDF文件转换为Word文档
    
    参数:
        pdf_path (str): PDF文件的路径
        output_path (str): 输出Word文档的路径，如果为None则自动生成
        
    返回:
        dict: 包含转换结果和文件路径的字典
    """
    if not DOCX_AVAILABLE:
        return {
            "status": "error",
            "message": "python-docx library is not available. Please install it to use Word conversion feature.",
            "word_path": None
        }
    
    doc = None
    try:
        # 验证文件路径
        print(f"尝试打开PDF文件: {pdf_path}")
        print(f"文件是否存在: {os.path.exists(pdf_path)}")
        
        if not os.path.exists(pdf_path):
            return {
                "status": "error",
                "message": f"文件不存在: {pdf_path}",
                "word_path": None
            }
        
        # 打开PDF文件
        doc = fitz.open(pdf_path)
        
        if doc.page_count == 0:
            return {
                "status": "error",
                "message": "PDF文件为空",
                "word_path": None
            }
        
        # 创建Word文档
        word_doc = Document()
        
        # 设置文档标题
        base_filename = os.path.splitext(os.path.basename(pdf_path))[0]
        title = word_doc.add_heading(f'{base_filename}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 创建图片存储目录
        images_dir = os.path.join(os.path.dirname(pdf_path), f"{base_filename}_images")
        if not os.path.exists(images_dir):
            os.makedirs(images_dir)
        
        # 提取文本块和格式信息
        text_elements = []
        extracted_images = []
        
        for page_num in range(doc.page_count):
            page = doc[page_num]
            
            # 添加页面分隔
            if page_num > 0:
                word_doc.add_page_break()
            
            # 添加页面标题
            page_heading = word_doc.add_heading(f'第 {page_num + 1} 页', level=1)
            page_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 提取图片
            page_images = extract_images_from_page(page, page_num, images_dir, base_filename)
            extracted_images.extend(page_images)
            
            # 获取文本字典，包含格式信息
            text_dict = page.get_text("dict")
            
            # 处理每个文本块
            page_text_elements = []
            for block in text_dict["blocks"]:
                if "lines" in block:  # 文本块
                    for line in block["lines"]:
                        for span in line["spans"]:
                            cleaned_text = clean_text_for_xml(span["text"])
                            if cleaned_text.strip():
                                page_text_elements.append({
                                    "text": cleaned_text,
                                    "font_size": span["size"],
                                    "font_flags": span["flags"],
                                    "bbox": span["bbox"],
                                    "page": page_num + 1
                                })
            
            # 将文本添加到Word文档
            add_text_to_word_doc(word_doc, page_text_elements)
            
            # 添加该页面的图片
            for img in page_images:
                try:
                    # 添加图片到Word文档
                    img_paragraph = word_doc.add_paragraph()
                    img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 检查图片文件是否存在
                    if os.path.exists(img['path']):
                        run = img_paragraph.add_run()
                        # 设置图片大小，最大宽度为6英寸
                        max_width = Inches(6)
                        run.add_picture(img['path'], width=max_width)
                        
                        # 添加图片说明
                        caption = word_doc.add_paragraph(f"图片 {img['index']} ({img['width']}x{img['height']})")
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_run = caption.runs[0]
                        caption_run.font.size = Pt(9)
                        caption_run.font.italic = True
                        
                except Exception as e:
                    print(f"添加图片到Word文档时出错: {str(e)}")
                    # 如果图片添加失败，添加文本说明
                    img_text = word_doc.add_paragraph(f"[图片: {img['filename']}]")
                    img_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 生成输出路径
        if output_path is None:
            output_path = os.path.join(os.path.dirname(pdf_path), f"{base_filename}.docx")
        
        # 保存Word文档
        word_doc.save(output_path)
        
        return {
            "status": "success",
            "message": "PDF成功转换为Word文档",
            "word_path": output_path,
            "images_count": len(extracted_images),
            "pages_count": doc.page_count if doc else 0
        }
    
    except Exception as e:
        # 记录错误并返回错误消息
        error_msg = f"转换过程中出错: {str(e)}"
        print(error_msg)
        return {
            "status": "error",
            "message": error_msg,
            "word_path": None
        }
    
    finally:
        # 确保PDF文档总是被关闭
        if doc is not None:
            try:
                doc.close()
            except:
                pass  # 忽略关闭时的错误

def clean_text_for_xml(text):
    """
    清理文本，移除不兼容XML的字符
    
    参数:
        text: 原始文本
    返回:
        清理后的文本
    """
    if not text:
        return ""
    
    # 移除NULL字节和控制字符
    import re
    # 保留常见的空白字符（空格、制表符、换行符），移除其他控制字符
    cleaned = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
    
    # 确保文本是有效的Unicode
    try:
        cleaned = cleaned.encode('utf-8', errors='ignore').decode('utf-8')
    except:
        cleaned = ""
    
    return cleaned

def add_text_to_word_doc(word_doc, text_elements):
    """
    将文本元素添加到Word文档中，保持格式
    
    参数:
        word_doc: Word文档对象
        text_elements: 文本元素列表
    """
    if not text_elements:
        return
    
    # 分析字体大小，确定标题级别
    font_sizes = [elem["font_size"] for elem in text_elements]
    avg_font_size = sum(font_sizes) / len(font_sizes)
    max_font_size = max(font_sizes)
    
    # 定义标题阈值
    title_threshold = avg_font_size * 1.2
    subtitle_threshold = avg_font_size * 1.1
    
    current_paragraph = None
    
    for i, elem in enumerate(text_elements):
        text = clean_text_for_xml(elem["text"]).strip()
        font_size = elem["font_size"]
        font_flags = elem["font_flags"]
        
        if not text:
            continue
        
        # 判断是否为标题
        if font_size >= title_threshold and font_size == max_font_size:
            # 主标题
            heading = word_doc.add_heading(text, level=2)
            current_paragraph = None
        elif font_size >= subtitle_threshold:
            # 副标题
            heading = word_doc.add_heading(text, level=3)
            current_paragraph = None
        else:
            # 普通文本
            if current_paragraph is None:
                current_paragraph = word_doc.add_paragraph()
            
            run = current_paragraph.add_run(text)
            
            # 设置字体格式
            if font_flags & 16:  # 粗体
                run.font.bold = True
            if font_flags & 2:   # 斜体
                run.font.italic = True
            
            # 设置字体大小
            run.font.size = Pt(max(9, min(font_size, 16)))
            
            # 检查是否需要换行
            next_elem = text_elements[i + 1] if i + 1 < len(text_elements) else None
            if next_elem:
                # 如果下一个元素的Y坐标差异较大，说明需要换行
                current_y = elem["bbox"][1]
                next_y = next_elem["bbox"][1]
                if abs(next_y - current_y) > font_size * 0.5:
                    current_paragraph = None
            else:
                current_paragraph = None

def pdf_to_cropped_images(pdf_path, output_dir=None):
    """
    将PDF文件的每一页裁剪为3:4比例的图片并打包为ZIP文件
    
    参数:
        pdf_path (str): PDF文件的路径
        output_dir (str): 输出目录，如果为None则使用PDF文件所在目录
        
    返回:
        dict: 包含转换结果和ZIP文件路径的字典
    """
    doc = None
    try:
        # 验证文件路径
        print(f"尝试打开PDF文件进行裁剪: {pdf_path}")
        print(f"文件是否存在: {os.path.exists(pdf_path)}")
        
        if not os.path.exists(pdf_path):
            return {
                "status": "error",
                "message": f"文件不存在: {pdf_path}",
                "zip_path": None
            }
        
        # 打开PDF文件
        doc = fitz.open(pdf_path)
        
        if doc.page_count == 0:
            return {
                "status": "error",
                "message": "PDF文件为空",
                "zip_path": None
            }
        
        # 生成输出路径
        base_filename = os.path.splitext(os.path.basename(pdf_path))[0]
        if output_dir is None:
            output_dir = os.path.dirname(pdf_path)
        
        # 创建临时目录存储图片
        temp_dir = os.path.join(output_dir, f"{base_filename}_cropped_images")
        os.makedirs(temp_dir, exist_ok=True)
        
        # 生成ZIP文件路径
        zip_path = os.path.join(output_dir, f"{base_filename}_cropped_images.zip")
        
        # 处理每一页
        cropped_images = []
        for page_num in range(doc.page_count):
            page = doc[page_num]
            
            # 获取页面尺寸
            page_rect = page.rect
            page_width = page_rect.width
            page_height = page_rect.height
            
            # 计算3:4比例的裁剪区域
            target_ratio = 3.0 / 4.0  # 宽:高 = 3:4
            current_ratio = page_width / page_height
            
            if current_ratio > target_ratio:
                # 页面太宽，需要裁剪宽度
                new_width = page_height * target_ratio
                x_offset = (page_width - new_width) / 2
                crop_rect = fitz.Rect(x_offset, 0, x_offset + new_width, page_height)
            else:
                # 页面太高，需要裁剪高度
                new_height = page_width / target_ratio
                y_offset = (page_height - new_height) / 2
                crop_rect = fitz.Rect(0, y_offset, page_width, y_offset + new_height)
            
            # 设置渲染参数，提高图片质量
            mat = fitz.Matrix(2.0, 2.0)  # 2倍缩放提高清晰度
            
            # 渲染裁剪后的页面为图片
            pix = page.get_pixmap(matrix=mat, clip=crop_rect)
            
            # 转换为PIL Image
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            
            # 确保图片比例为3:4
            img_width, img_height = img.size
            target_width = int(img_height * target_ratio)
            if img_width != target_width:
                # 微调宽度以确保精确的3:4比例
                left = (img_width - target_width) // 2
                right = left + target_width
                img = img.crop((left, 0, right, img_height))
            
            # 保存图片
            img_filename = f"page_{page_num + 1:03d}.png"
            img_path = os.path.join(temp_dir, img_filename)
            img.save(img_path, "PNG", quality=95)
            
            cropped_images.append({
                "page": page_num + 1,
                "filename": img_filename,
                "path": img_path,
                "size": img.size
            })
            
            print(f"已处理第 {page_num + 1} 页，图片尺寸: {img.size}")
        
        # 创建ZIP文件
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for img_info in cropped_images:
                zipf.write(img_info["path"], img_info["filename"])
        
        # 清理临时文件
        for img_info in cropped_images:
            try:
                os.remove(img_info["path"])
            except:
                pass
        
        # 删除临时目录
        try:
            os.rmdir(temp_dir)
        except:
            pass
        
        return {
            "status": "success",
            "message": f"PDF成功裁剪为{len(cropped_images)}张3:4比例图片",
            "zip_path": zip_path,
            "images_count": len(cropped_images),
            "pages_count": doc.page_count
        }
    
    except Exception as e:
        # 记录错误并返回错误消息
        error_msg = f"裁剪过程中出错: {str(e)}"
        print(error_msg)
        return {
            "status": "error",
            "message": error_msg,
            "zip_path": None
        }
    
    finally:
        # 确保PDF文档总是被关闭
        if doc is not None:
            try:
                doc.close()
            except:
                pass  # 忽略关闭时的错误